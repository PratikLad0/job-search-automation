"""
DOCX Generator for Resumes and Cover Letters.
Uses python-docx to create editable Word documents.
"""

import logging
from pathlib import Path
from typing import Optional, Dict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.app.core import config
from backend.app.db.models import Job, CVData
from backend.app.services.parsers.cv_parser import parse_cv
from backend.app.services.ai.resume_tailor import tailor_resume
from backend.app.services.ai.cover_letter import generate_cover_letter

logger = logging.getLogger(__name__)

def _set_document_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

def _add_heading_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(24)
        run.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80) # Dark Blue
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        # Add bottom border
        pPr = heading._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

def generate_resume_docx(job: Job, cv_data: Optional[CVData] = None) -> Optional[Path]:
    if cv_data is None:
        cv_data = parse_cv()

    try:
        tailored = tailor_resume(job, cv_data)
        doc = Document()
        _set_document_margins(doc)

        # Header
        _add_heading_style(doc, cv_data.name, 1)
        
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if cv_data.email: contact_info.append(cv_data.email)
        if cv_data.phone: contact_info.append(cv_data.phone)
        if cv_data.location: contact_info.append(cv_data.location)
        if cv_data.linkedin_url: contact_info.append("LinkedIn")
        
        contact.add_run(" | ".join(contact_info))
        
        # Summary
        _add_heading_style(doc, "Professional Summary", 2)
        doc.add_paragraph(tailored.get("summary", cv_data.summary))

        # Skills
        _add_heading_style(doc, "Skills", 2)
        skills_text = ", ".join(tailored.get("skills_highlighted", cv_data.skills))
        doc.add_paragraph(skills_text)

        # Experience
        _add_heading_style(doc, "Experience", 2)
        expenses = tailored.get("experience", cv_data.experience)
        for exp in expenses:
            p = doc.add_paragraph()
            title_run = p.add_run(exp.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(11)
            p.add_run(f" | {exp.get('company', '')}").italic = True
            
            # Right align date if possible, but for simplicity just append
            p.add_run(f" | {exp.get('duration', '')}")

            bullets = exp.get("bullets", [])
            # Fallback if AI didn't return bullets but description text
            if not bullets and exp.get("description"):
                doc.add_paragraph(exp.get("description")).style = 'List Bullet'
            
            for bullet in bullets:
                doc.add_paragraph(bullet).style = 'List Bullet'

        # Projects (Highlighted)
        projects = tailored.get("projects_highlighted", cv_data.projects[:3])
        if projects:
            _add_heading_style(doc, "Key Projects", 2)
            for proj in projects:
                p = doc.add_paragraph()
                name_run = p.add_run(proj.get("name", ""))
                name_run.bold = True
                p.add_run(" - " + proj.get("description", ""))
                if proj.get("tech_stack"):
                    p.add_run(f" (Tech: {', '.join(proj.get('tech_stack'))})").italic = True

        # Education
        _add_heading_style(doc, "Education", 2)
        education = tailored.get("education", cv_data.education)
        for edu in education:
            p = doc.add_paragraph()
            p.add_run(edu.get("degree", "")).bold = True
            p.add_run(f", {edu.get('institution', '')}")
            p.add_run(f" ({edu.get('year', '')})")

        # Save
        safe_company = "".join(x for x in job.company if x.isalnum())
        safe_title = "".join(x for x in job.title if x.isalnum())
        filename = f"resume_{safe_company}_{safe_title}.docx"
        output_path = config.RESUMES_DIR / filename
        doc.save(output_path)
        logger.info(f"DOCX Resume generated: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return None

def generate_cover_letter_docx(job: Job, cv_data: Optional[CVData] = None) -> Optional[Path]:
    if cv_data is None:
        cv_data = parse_cv()

    try:
        letter_text = generate_cover_letter(job, cv_data)
        doc = Document()
        _set_document_margins(doc)

        # Header
        _add_heading_style(doc, cv_data.name, 1)
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"{cv_data.email} | {cv_data.phone}")

        doc.add_paragraph() # Spacer

        # Body
        for para in letter_text.split("\n\n"):
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)

        # Save
        safe_company = "".join(x for x in job.company if x.isalnum())
        safe_title = "".join(x for x in job.title if x.isalnum())
        filename = f"cover_letter_{safe_company}_{safe_title}.docx"
        output_path = config.COVER_LETTERS_DIR / filename
        doc.save(output_path)
        logger.info(f"DOCX Cover Letter generated: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"DOCX Cover Letter generation failed: {e}")
        return None
